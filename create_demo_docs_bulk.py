"""
Generate demo .docx documents for DocuAlign AI demonstration.
Creates 10 diverse documents with intentional contradictions for testing.

Run: python create_demo_docs_bulk.py
Output: demo_docs/ folder with 10 .docx files
"""
import os
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from PIL import Image, ImageDraw, ImageFont

DEMO_DIR = Path("demo_docs")
IMG_DIR = DEMO_DIR / "images"


def create_dirs():
    DEMO_DIR.mkdir(exist_ok=True)
    IMG_DIR.mkdir(exist_ok=True)


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def _draw_rounded_rect(draw, xy, radius, fill):
    x0, y0, x1, y1 = xy
    draw.rectangle([x0 + radius, y0, x1 - radius, y1], fill=fill)
    draw.rectangle([x0, y0 + radius, x1, y1 - radius], fill=fill)
    draw.pieslice([x0, y0, x0 + 2*radius, y0 + 2*radius], 180, 270, fill=fill)
    draw.pieslice([x1 - 2*radius, y0, x1, y0 + 2*radius], 270, 360, fill=fill)
    draw.pieslice([x0, y1 - 2*radius, x0 + 2*radius, y1], 90, 180, fill=fill)
    draw.pieslice([x1 - 2*radius, y1 - 2*radius, x1, y1], 0, 90, fill=fill)


def _add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def _add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def create_simple_screenshot(filename, title_text, bg_color="#f0f2f5"):
    """Create a simple placeholder screenshot."""
    img = Image.new("RGB", (800, 500), bg_color)
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, 800, 60], fill="#1a73e8")
    draw.text((20, 18), title_text, fill="white")
    path = IMG_DIR / filename
    img.save(path)
    return path


# ---------------------------------------------------------------------------
# Document templates
# ---------------------------------------------------------------------------

def create_document_1():
    """API Reference Guide v2.0 (outdated endpoint info)."""
    doc = Document()
    doc.add_heading("API Reference Guide v2.0", level=0)
    _add_para(doc, "最終更新: 2024年5月1日")
    doc.add_paragraph()
    
    _add_heading(doc, "Authentication Endpoint", 1)
    _add_para(doc, "認証エンドポイント: POST /api/v2/auth")
    _add_para(doc, "リクエストボディに username, password を含めてください。")
    
    _add_heading(doc, "User Management", 1)
    _add_para(doc, "ユーザー一覧取得: GET /api/v2/users")
    _add_para(doc, "ページネーション最大: 50件")
    
    path = DEMO_DIR / "API_Reference_Guide_v2.0.docx"
    doc.save(str(path))
    print(f"[OK] Created: {path}")


def create_document_2():
    """API Migration Guide v3.0 (new info contradicting v2.0)."""
    doc = Document()
    doc.add_heading("API Migration Guide v3.0", level=0)
    _add_para(doc, "公開日: 2025年2月1日")
    doc.add_paragraph()
    
    _add_heading(doc, "Breaking Changes", 1)
    _add_para(doc, "認証エンドポイントが /api/v3/authenticate に変更されました。")
    _add_para(doc, "OAuth 2.0 を採用したため、username/password 方式は廃止されました。")
    
    _add_heading(doc, "Pagination Update", 1)
    _add_para(doc, "ページネーション最大件数を 50件 → 100件 に拡大しました。")
    
    path = DEMO_DIR / "API_Migration_Guide_v3.0.docx"
    doc.save(str(path))
    print(f"[OK] Created: {path}")


def create_document_3():
    """Employee Handbook 2024 (with outdated policy)."""
    doc = Document()
    doc.add_heading("社員ハンドブック 2024年版", level=0)
    _add_para(doc, "人事部発行")
    doc.add_paragraph()
    
    _add_heading(doc, "勤務時間", 1)
    _add_para(doc, "コアタイム: 10:00-15:00")
    _add_para(doc, "フレックス: 8:00-20:00 の範囲で自由設定")
    
    _add_heading(doc, "有給休暇", 1)
    _add_para(doc, "年次有給休暇: 入社年度は10日、翌年度以降は20日")
    _add_para(doc, "申請は Slack の #hr-request チャンネルで行ってください。")
    
    path = DEMO_DIR / "Employee_Handbook_2024.docx"
    doc.save(str(path))
    print(f"[OK] Created: {path}")


def create_document_4():
    """HR Policy Update 2025 (contradicts handbook)."""
    doc = Document()
    doc.add_heading("人事制度改定のお知らせ 2025", level=0)
    _add_para(doc, "2025年4月1日施行")
    doc.add_paragraph()
    
    _add_heading(doc, "フレックスタイム制の変更", 1)
    _add_para(doc, "コアタイムを廃止し、完全フレックスに移行します。")
    
    _add_heading(doc, "有給申請方法の変更", 1)
    _add_para(doc, "有給申請は新システム「HR Portal」から行ってください。")
    _add_para(doc, "Slack での申請は3月31日で終了しました。")
    
    path = DEMO_DIR / "HR_Policy_Update_2025.docx"
    doc.save(str(path))
    print(f"[OK] Created: {path}")


def create_document_5():
    """Security Guidelines v1.5 (password policy)."""
    doc = Document()
    doc.add_heading("セキュリティガイドライン v1.5", level=0)
    _add_para(doc, "情報セキュリティ部")
    doc.add_paragraph()
    
    _add_heading(doc, "パスワードポリシー", 1)
    _add_para(doc, "最低文字数: 8文字")
    _add_para(doc, "英数字混在必須")
    _add_para(doc, "3ヶ月ごとに変更してください")
    
    _add_heading(doc, "ファイル共有", 1)
    _add_para(doc, "社外とのファイル共有は Dropbox を使用してください。")
    
    path = DEMO_DIR / "Security_Guidelines_v1.5.docx"
    doc.save(str(path))
    print(f"[OK] Created: {path}")


def create_document_6():
    """Security Policy Update v2.0 (contradicts v1.5)."""
    doc = Document()
    doc.add_heading("セキュリティポリシー v2.0", level=0)
    _add_para(doc, "2025年3月施行")
    doc.add_paragraph()
    
    _add_heading(doc, "パスワードポリシー改定", 1)
    _add_para(doc, "最低文字数: 12文字に引き上げ")
    _add_para(doc, "記号を1文字以上含めることを必須化")
    _add_para(doc, "定期変更は廃止（強度の高いパスワードを継続使用）")
    
    _add_heading(doc, "ファイル共有ツール変更", 1)
    _add_para(doc, "社外共有は Google Drive のみ許可。Dropbox は使用禁止です。")
    
    path = DEMO_DIR / "Security_Policy_v2.0.docx"
    doc.save(str(path))
    print(f"[OK] Created: {path}")


def create_document_7():
    """IT Support FAQ (outdated contact info)."""
    doc = Document()
    doc.add_heading("IT サポート FAQ", level=0)
    _add_para(doc, "最終更新: 2024年1月")
    doc.add_paragraph()
    
    _add_heading(doc, "Q. パスワードを忘れました", 1)
    _add_para(doc, "A. IT部門（内線: 1234）に電話してリセット依頼をしてください。")
    
    _add_heading(doc, "Q. PCが起動しません", 1)
    _add_para(doc, "A. IT部門（内線: 1234）までご連絡ください。")
    
    path = DEMO_DIR / "IT_Support_FAQ.docx"
    doc.save(str(path))
    print(f"[OK] Created: {path}")


def create_document_8():
    """IT Department Announcement (new contact method)."""
    doc = Document()
    doc.add_heading("IT部門 連絡先変更のお知らせ", level=0)
    _add_para(doc, "2025年1月15日から")
    doc.add_paragraph()
    
    _add_heading(doc, "新しい連絡方法", 1)
    _add_para(doc, "内線1234は廃止されました。")
    _add_para(doc, "今後のお問い合わせは Slack #it-support チャンネルでお願いします。")
    _add_para(doc, "緊急時は ticketシステム（https://ticket.example.com）をご利用ください。")
    
    path = DEMO_DIR / "IT_Contact_Update.docx"
    doc.save(str(path))
    print(f"[OK] Created: {path}")


def create_document_9():
    """Remote Work Policy (initial version)."""
    doc = Document()
    doc.add_heading("リモートワーク規定", level=0)
    _add_para(doc, "2024年4月制定")
    doc.add_paragraph()
    
    _add_heading(doc, "リモートワーク可能日数", 1)
    _add_para(doc, "週2日まで")
    
    _add_heading(doc, "申請方法", 1)
    _add_para(doc, "前日までにマネージャーにメールで申請")
    
    _add_heading(doc, "勤怠管理", 1)
    _add_para(doc, "始業・終業時に勤怠システムに打刻してください")
    
    path = DEMO_DIR / "Remote_Work_Policy.docx"
    doc.save(str(path))
    print(f"[OK] Created: {path}")


def create_document_10():
    """Remote Work Policy Update (contradicts initial)."""
    doc = Document()
    doc.add_heading("リモートワーク規定 改定版", level=0)
    _add_para(doc, "2025年2月改定")
    doc.add_paragraph()
    
    _add_heading(doc, "リモートワーク日数制限撤廃", 1)
    _add_para(doc, "週2日の制限を撤廃し、フルリモート勤務も可能になりました。")
    
    _add_heading(doc, "申請方法の変更", 1)
    _add_para(doc, "メール申請は廃止。HR Portal から事前登録してください。")
    
    _add_heading(doc, "勤怠管理の簡素化", 1)
    _add_para(doc, "リモート時の打刻は不要になりました（自動記録）。")
    
    path = DEMO_DIR / "Remote_Work_Policy_Update.docx"
    doc.save(str(path))
    print(f"[OK] Created: {path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    create_dirs()
    
    print(">> 10件のデモドキュメントを生成中...\n")
    
    create_document_1()
    create_document_2()
    create_document_3()
    create_document_4()
    create_document_5()
    create_document_6()
    create_document_7()
    create_document_8()
    create_document_9()
    create_document_10()
    
    print("\n>> 完了! demo_docs/ フォルダに10件のファイルが作成されました")
    print("   生成ファイル:")
    for f in sorted(DEMO_DIR.glob("*.docx")):
        print(f"   - {f.name}")
