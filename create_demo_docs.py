"""
Generate demo .docx documents for DocuGardener Agent demonstration.
Creates:
  1. Old manual (v2.1) with outdated screenshots and instructions
  2. New release notes (v3.0) with updated information
  3. Onboarding guide with stale information

Run: python create_demo_docs.py
Output: demo_docs/ folder with .docx files
"""
import os
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from PIL import Image, ImageDraw, ImageFont

DEMO_DIR = Path("demo_docs")
IMG_DIR = DEMO_DIR / "images"


def create_dirs():
    DEMO_DIR.mkdir(exist_ok=True)
    IMG_DIR.mkdir(exist_ok=True)


# ---------------------------------------------------------------------------
# Screenshot generation (Pillow)
# ---------------------------------------------------------------------------

def _draw_rounded_rect(draw, xy, radius, fill):
    x0, y0, x1, y1 = xy
    draw.rectangle([x0 + radius, y0, x1 - radius, y1], fill=fill)
    draw.rectangle([x0, y0 + radius, x1, y1 - radius], fill=fill)
    draw.pieslice([x0, y0, x0 + 2*radius, y0 + 2*radius], 180, 270, fill=fill)
    draw.pieslice([x1 - 2*radius, y0, x1, y0 + 2*radius], 270, 360, fill=fill)
    draw.pieslice([x0, y1 - 2*radius, x0 + 2*radius, y1], 90, 180, fill=fill)
    draw.pieslice([x1 - 2*radius, y1 - 2*radius, x1, y1], 0, 90, fill=fill)


def create_old_login_screenshot():
    """Create an 'old' login page screenshot (v2.1 style)."""
    img = Image.new("RGB", (800, 500), "#f0f2f5")
    draw = ImageDraw.Draw(img)

    # Header bar (blue)
    draw.rectangle([0, 0, 800, 60], fill="#1a73e8")
    draw.text((20, 18), "社内ポータル v2.1", fill="white")
    # Gear icon area (top right)
    draw.ellipse([740, 15, 770, 45], fill="#ffffff", outline="#ffffff")
    draw.text((748, 20), "⚙", fill="#1a73e8")

    # Login form
    _draw_rounded_rect(draw, (200, 120, 600, 420), 12, "#ffffff")
    draw.text((320, 140), "ログイン", fill="#333333")

    # Fields
    draw.text((230, 190), "ユーザーID", fill="#666666")
    _draw_rounded_rect(draw, (230, 210, 570, 245), 6, "#f5f5f5")
    draw.rectangle([230, 210, 570, 245], outline="#cccccc")

    draw.text((230, 265), "パスワード", fill="#666666")
    _draw_rounded_rect(draw, (230, 285, 570, 320), 6, "#f5f5f5")
    draw.rectangle([230, 285, 570, 320], outline="#cccccc")

    # Login button
    _draw_rounded_rect(draw, (230, 350, 570, 390), 8, "#1a73e8")
    draw.text((370, 360), "ログイン", fill="#ffffff")

    path = IMG_DIR / "old_login.png"
    img.save(path)
    return path


def create_new_login_screenshot():
    """Create a 'new' login page screenshot (v3.0 style with SSO)."""
    img = Image.new("RGB", (800, 500), "#0f1117")
    draw = ImageDraw.Draw(img)

    # Sidebar (new design - no gear icon here)
    draw.rectangle([0, 0, 200, 500], fill="#1a1d29")
    draw.text((20, 20), "社内ポータル v3.0", fill="#22c55e")
    draw.text((20, 60), "▸ ホーム", fill="#94a3b8")
    draw.text((20, 90), "▸ 通知センター", fill="#94a3b8")
    draw.text((20, 120), "▸ ファイル共有", fill="#94a3b8")
    draw.text((20, 150), "▸ チャット", fill="#94a3b8")
    draw.text((20, 180), "▸ 設定", fill="#22c55e")

    # Login form (modern dark theme)
    _draw_rounded_rect(draw, (300, 80, 700, 430), 16, "#1a1d29")
    draw.text((430, 105), "ログイン", fill="#e2e8f0")

    # SSO button
    _draw_rounded_rect(draw, (330, 160, 670, 200), 8, "#4285f4")
    draw.text((390, 170), "Google SSO でログイン", fill="#ffffff")

    draw.text((470, 220), "または", fill="#94a3b8")

    # Fields
    draw.text((330, 255), "メールアドレス", fill="#94a3b8")
    _draw_rounded_rect(draw, (330, 275, 670, 310), 6, "#222639")
    draw.rectangle([330, 275, 670, 310], outline="#2d3348")

    draw.text((330, 330), "パスワード", fill="#94a3b8")
    _draw_rounded_rect(draw, (330, 350, 670, 385), 6, "#222639")
    draw.rectangle([330, 350, 670, 385], outline="#2d3348")

    # Login button
    _draw_rounded_rect(draw, (330, 400, 670, 435), 8, "#22c55e")
    draw.text((450, 408), "ログイン", fill="#ffffff")

    path = IMG_DIR / "new_login.png"
    img.save(path)
    return path


def create_old_settings_screenshot():
    """Create an 'old' settings page screenshot (gear icon in top right)."""
    img = Image.new("RGB", (800, 500), "#f8f9fa")
    draw = ImageDraw.Draw(img)

    # Header
    draw.rectangle([0, 0, 800, 60], fill="#1a73e8")
    draw.text((20, 18), "社内ポータル v2.1 > 設定", fill="white")
    draw.ellipse([740, 15, 770, 45], fill="#ffffff")
    draw.text((748, 20), "⚙", fill="#1a73e8")

    # Settings panel
    _draw_rounded_rect(draw, (20, 80, 780, 460), 8, "#ffffff")
    draw.text((40, 100), "設定画面", fill="#333333")
    draw.line([40, 130, 760, 130], fill="#eeeeee", width=1)

    items = ["通知設定", "言語設定", "プロフィール編集", "セキュリティ", "表示設定"]
    for i, item in enumerate(items):
        y = 150 + i * 55
        _draw_rounded_rect(draw, (40, y, 760, y + 40), 6, "#f5f7fa")
        draw.text((60, y + 10), f"▸ {item}", fill="#333333")

    path = IMG_DIR / "old_settings.png"
    img.save(path)
    return path


def create_old_dashboard_screenshot():
    """Create an 'old' dashboard screenshot."""
    img = Image.new("RGB", (800, 500), "#f8f9fa")
    draw = ImageDraw.Draw(img)

    # Header
    draw.rectangle([0, 0, 800, 60], fill="#1a73e8")
    draw.text((20, 18), "社内ポータル v2.1 > ダッシュボード", fill="white")
    draw.ellipse([740, 15, 770, 45], fill="#ffffff")
    draw.text((748, 20), "⚙", fill="#1a73e8")

    # Left menu
    draw.rectangle([0, 60, 180, 500], fill="#ffffff")
    menus = ["ダッシュボード", "マイドライブ", "共有フォルダ", "コミュニケーション", "お知らせ"]
    for i, m in enumerate(menus):
        y = 80 + i * 40
        if i == 0:
            draw.rectangle([0, y, 180, y + 35], fill="#e8f0fe")
        draw.text((15, y + 8), m, fill="#333333")

    # Dashboard cards
    cards = [("未読通知", "12"), ("共有ファイル", "48"), ("チャット", "3")]
    for i, (label, count) in enumerate(cards):
        x = 200 + i * 195
        _draw_rounded_rect(draw, (x, 80, x + 180, 180), 8, "#ffffff")
        draw.rectangle([x, 80, x + 180, 180], outline="#e0e0e0")
        draw.text((x + 70, 100), count, fill="#1a73e8")
        draw.text((x + 40, 145), label, fill="#666666")

    path = IMG_DIR / "old_dashboard.png"
    img.save(path)
    return path


# ---------------------------------------------------------------------------
# Document creation
# ---------------------------------------------------------------------------

def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def _add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def create_old_manual(login_img, settings_img, dashboard_img):
    """Create the old operations manual (v2.1) .docx."""
    doc = Document()

    # Title
    title = doc.add_heading("社内ポータル操作手順書 v2.1", level=0)
    _add_para(doc, "最終更新: 2024年4月15日")
    _add_para(doc, "管理部門: 情報システム部")
    _add_para(doc, "対象者: 全社員")
    doc.add_paragraph()

    # TOC
    _add_heading(doc, "目次", 1)
    toc_items = [
        "1. ログイン方法",
        "2. ダッシュボードの使い方",
        "3. 設定画面の開き方",
        "4. ファイル共有",
        "5. チャット機能",
        "6. お知らせ確認",
        "7. トラブルシューティング",
    ]
    for item in toc_items:
        _add_para(doc, item)
    doc.add_page_break()

    # Section 1: Login
    _add_heading(doc, "1. ログイン方法", 1)
    _add_para(doc, "ブラウザで https://portal.example.com にアクセスし、社員IDとパスワードを入力してログインします。")
    _add_para(doc, "")
    _add_para(doc, "【ログイン画面】", bold=True)
    doc.add_picture(str(login_img), width=Inches(5.5))
    _add_para(doc, "図1: ログイン画面（ユーザーIDとパスワードを入力）")
    doc.add_paragraph()
    _add_para(doc, "※ パスワードを忘れた場合は、IT部門（内線: 1234）にご連絡ください。")

    # Section 2: Dashboard
    _add_heading(doc, "2. ダッシュボードの使い方", 1)
    _add_para(doc, "ログイン後、ダッシュボードが表示されます。左側に主要メニュー、右側に通知パネルがあります。")
    _add_para(doc, "")
    _add_para(doc, "【ダッシュボード画面】", bold=True)
    doc.add_picture(str(dashboard_img), width=Inches(5.5))
    _add_para(doc, "図2: ダッシュボード（左メニューから各機能にアクセス）")

    # Section 3: Settings
    _add_heading(doc, "3. 設定画面の開き方", 1)
    _add_para(doc, "設定画面は右上のギアアイコン（⚙）から開きます。通知設定、言語設定、プロフィール編集が行えます。")
    _add_para(doc, "")
    _add_para(doc, "【設定画面へのアクセス】", bold=True)
    doc.add_picture(str(settings_img), width=Inches(5.5))
    _add_para(doc, "図3: 右上のギアアイコンをクリックして設定画面を開く")
    doc.add_paragraph()
    _add_para(doc, "手順:")
    steps = [
        "1. 画面右上のギアアイコン（⚙）をクリック",
        "2. ドロップダウンメニューから「設定」を選択",
        "3. 各設定項目を編集",
        "4. 「保存」ボタンをクリック",
    ]
    for step in steps:
        _add_para(doc, step)

    # Section 4: File sharing
    _add_heading(doc, "4. ファイル共有", 1)
    _add_para(doc, "共有フォルダは「マイドライブ」>「共有」から開きます。")
    _add_para(doc, "アップロード上限は100MBです。100MBを超えるファイルは分割してアップロードしてください。")

    # Section 5: Chat
    _add_heading(doc, "5. チャット機能", 1)
    _add_para(doc, "社内チャットは左メニューの「コミュニケーション」から利用できます。")
    _add_para(doc, "グループチャットは最大5人まで作成可能です。6人以上の場合はメーリングリストをご利用ください。")

    # Section 6: Announcements
    _add_heading(doc, "6. お知らせ確認", 1)
    _add_para(doc, "お知らせは左メニューの「お知らせ」から確認できます。重要なお知らせは赤いバッジで表示されます。")

    # Section 7: Troubleshooting
    _add_heading(doc, "7. トラブルシューティング", 1)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "症状"
    hdr_cells[1].text = "対処法"
    data = [
        ("ログインできない", "パスワードリセット（IT部門: 内線1234）"),
        ("画面が表示されない", "ブラウザキャッシュをクリア"),
        ("ファイルがアップロードできない", "100MB以下か確認。分割アップロード推奨"),
        ("チャットが送信できない", "ネットワーク接続を確認"),
    ]
    for i, (symptom, fix) in enumerate(data):
        row = table.rows[i + 1]
        row.cells[0].text = symptom
        row.cells[1].text = fix

    path = DEMO_DIR / "社内ポータル操作手順書_v2.1.docx"
    doc.save(str(path))
    print(f"✅ Created: {path}")
    return path


def create_new_release_notes():
    """Create the new release notes (v3.0) .docx."""
    doc = Document()

    title = doc.add_heading("社内ポータル v3.0 リリースノート", level=0)
    _add_para(doc, "公開日: 2025年1月10日")
    _add_para(doc, "発行: 情報システム部")
    doc.add_paragraph()

    _add_heading(doc, "主な変更点", 1)

    changes = [
        ("UIの刷新", "ダークモードベースの新デザインに移行。ヘッダーバーを廃止し、サイドメニューに統合。"),
        ("設定画面の移動", "設定画面はサイドメニューに移動しました。右上のギアアイコンは廃止されました。"),
        ("名称変更", "「ダッシュボード」は「ホーム画面」に名称変更しました。"),
        ("ファイル共有の改善", "アップロード上限を100MBから500MBに拡大しました。"),
        ("チャット機能の強化", "グループチャットの人数制限を撤廃（5人→無制限）。音声通話機能を追加。"),
        ("通知センター新設", "全通知を一元管理する「通知センター」をサイドメニューに新設。"),
        ("SSO対応", "Google SSOによるシングルサインオンに対応。従来のID/パスワード認証も継続利用可。"),
        ("IT部門連絡先変更", "IT部門の連絡先がSlackチャンネル #it-support に変更（内線1234は廃止）。"),
    ]

    for title_text, desc in changes:
        _add_heading(doc, title_text, 2)
        _add_para(doc, desc)

    doc.add_page_break()
    _add_heading(doc, "影響を受けるドキュメント", 1)
    _add_para(doc, "以下のドキュメントの更新が必要です:")
    affected = [
        "社内ポータル操作手順書 v2.1 → 全セクション要更新",
        "新入社員向けガイド 2024年版 → セクション2, 3 要更新",
        "IT部門FAQ集 → 連絡先情報の更新",
    ]
    for item in affected:
        p = doc.add_paragraph(item)
        p.style = "List Bullet"

    path = DEMO_DIR / "社内ポータル_v3.0_リリースノート.docx"
    doc.save(str(path))
    print(f"✅ Created: {path}")
    return path


def create_onboarding_guide(dashboard_img):
    """Create the onboarding guide .docx with stale info."""
    doc = Document()

    doc.add_heading("新入社員向けガイド 2024年版", level=0)
    _add_para(doc, "最終更新: 2024年3月1日")
    _add_para(doc, "人事部 / 情報システム部")
    doc.add_paragraph()

    _add_heading(doc, "1. 初日にやること", 1)
    checklist = [
        "□ 社員証を受け取る（総務部 3F）",
        "□ PCセットアップ（IT部門が対応 / 内線: 1234）",
        "□ 社内ポータルにログイン",
        "□ 部門チャットグループに参加",
        "□ 社内研修動画を視聴",
    ]
    for item in checklist:
        _add_para(doc, item)

    _add_heading(doc, "2. 社内ポータルの使い方", 1)
    _add_para(doc, "ログイン後、ダッシュボードから各機能にアクセスできます。")
    _add_para(doc, "右上のギアアイコン（⚙）で通知やプロフィールの設定を変更できます。")
    _add_para(doc, "")
    _add_para(doc, "【ダッシュボード画面】", bold=True)
    doc.add_picture(str(dashboard_img), width=Inches(5.0))
    _add_para(doc, "図: ダッシュボード（ログイン直後の画面）")

    _add_heading(doc, "3. コミュニケーション", 1)
    _add_para(doc, "社内チャットでチームメンバーと連絡が取れます。")
    _add_para(doc, "グループチャットは最大5人まで作成可能です。それ以上の場合はメーリングリストを利用してください。")

    _add_heading(doc, "4. ファイル共有", 1)
    _add_para(doc, "「マイドライブ」>「共有」からファイルを共有できます。")
    _add_para(doc, "アップロード上限は100MBです。")

    _add_heading(doc, "5. トラブル時の連絡先", 1)
    _add_para(doc, "IT部門: 内線 1234")
    _add_para(doc, "総務部: 内線 5678")
    _add_para(doc, "人事部: 内線 9012")

    path = DEMO_DIR / "新入社員向けガイド_2024年版.docx"
    doc.save(str(path))
    print(f"✅ Created: {path}")
    return path


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    create_dirs()

    # Generate screenshots
    print("📸 スクリーンショット生成中...")
    old_login = create_old_login_screenshot()
    new_login = create_new_login_screenshot()
    old_settings = create_old_settings_screenshot()
    old_dashboard = create_old_dashboard_screenshot()
    print(f"   → {old_login}, {new_login}, {old_settings}, {old_dashboard}")

    # Generate .docx files
    print("\n📄 .docx ファイル生成中...")
    create_old_manual(old_login, old_settings, old_dashboard)
    create_new_release_notes()
    create_onboarding_guide(old_dashboard)

    print("\n🌿 完了！ demo_docs/ フォルダを確認してください")
    print("   生成ファイル:")
    for f in DEMO_DIR.glob("*.docx"):
        print(f"   - {f.name}")
