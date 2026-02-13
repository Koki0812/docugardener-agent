"""
Generate comprehensive demo documents for all test scenarios.
Creates both auto-fixable (.docx, .txt, .md) and manual-action (.pdf simulation) files.

Run: python create_demo_scenarios.py
Output: demo_docs/auto_fix/ and demo_docs/manual_action/
"""
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt

# Directories
BASE_DIR = Path("demo_docs")
AUTO_FIX_DIR = BASE_DIR / "auto_fix"
MANUAL_DIR = BASE_DIR / "manual_action"

def create_dirs():
    AUTO_FIX_DIR.mkdir(parents=True, exist_ok=True)
    MANUAL_DIR.mkdir(parents=True, exist_ok=True)

def _add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p

# ============================================================================
# AUTO-FIXABLE DOCUMENTS
# ============================================================================

def create_a1_ui_guide_v2():
    """A-1: UI要素名称変更 (旧版)"""
    doc = Document()
    doc.add_heading("UI操作ガイド v2.0", level=0)
    _add_para(doc, "最終更新: 2024年6月1日")
    doc.add_paragraph()
    
    doc.add_heading("設定画面へのアクセス", level=1)
    _add_para(doc, "画面右上のギアアイコン（⚙）をクリックして設定画面を開きます。")
    _add_para(doc, "ここから通知設定、言語設定、プロフィール編集が可能です。")
    
    path = AUTO_FIX_DIR / "UI_Guide_v2.docx"
    doc.save(str(path))
    print(f"[OK] {path}")

def create_a1_ui_changelog_v3():
    """A-1: UI要素名称変更 (新版)"""
    doc = Document()
    doc.add_heading("UI変更ログ v3.0", level=0)
    _add_para(doc, "リリース日: 2025年2月1日")
    doc.add_paragraph()
    
    doc.add_heading("重要な変更", level=1)
    _add_para(doc, "設定画面は右上のギアアイコンから「サイドメニュー」に移動しました。")
    _add_para(doc, "左側のサイドメニューから「設定」を選択してください。")
    
    path = AUTO_FIX_DIR / "UI_Changelog_v3.docx"
    doc.save(str(path))
    print(f"[OK] {path}")

def create_a2_onboarding_2024():
    """A-2: 用語の不統一 (旧版)"""
    doc = Document()
    doc.add_heading("新入社員オンボーディング 2024", level=0)
    doc.add_paragraph()
    
    doc.add_heading("初日の手順", level=1)
    _add_para(doc, "ログイン後、ダッシュボードから各機能にアクセスできます。")
    _add_para(doc, "ダッシュボードには未読通知、共有ファイル、チャットが表示されます。")
    
    path = AUTO_FIX_DIR / "Onboarding_Guide_2024.docx"
    doc.save(str(path))
    print(f"[OK] {path}")

def create_a2_product_update_2025():
    """A-2: 用語の不統一 (新版)"""
    doc = Document()
    doc.add_heading("製品アップデート 2025", level=0)
    doc.add_paragraph()
    
    doc.add_heading("用語変更のお知らせ", level=1)
    _add_para(doc, "「ダッシュボード」は「ホーム画面」に名称変更されました。")
    _add_para(doc, "今後は「ホーム画面」という呼称に統一してください。")
    
    path = AUTO_FIX_DIR / "Product_Update_2025.docx"
    doc.save(str(path))
    print(f"[OK] {path}")

def create_a3_api_ref_v2():
    """A-3: APIエンドポイント変更 (旧版)"""
    doc = Document()
    doc.add_heading("API リファレンス v2.0", level=0)
    doc.add_paragraph()
    
    doc.add_heading("認証エンドポイント", level=1)
    _add_para(doc, "POST /api/v2/auth")
    _add_para(doc, "リクエストボディ: { username, password }")
    
    path = AUTO_FIX_DIR / "API_Reference_v2.docx"
    doc.save(str(path))
    print(f"[OK] {path}")

def create_a3_api_migration_v3():
    """A-3: APIエンドポイント変更 (新版)"""
    doc = Document()
    doc.add_heading("API Migration Guide v3.0", level=0)
    doc.add_paragraph()
    
    doc.add_heading("Breaking Changes", level=1)
    _add_para(doc, "認証エンドポイントが POST /api/v3/authenticate に変更されました。")
    _add_para(doc, "OAuth 2.0 を採用したため、username/password 形式は廃止されました。")
    
    path = AUTO_FIX_DIR / "API_Migration_v3.docx"
    doc.save(str(path))
    print(f"[OK] {path}")

def create_a4_it_faq():
    """A-4: 連絡先情報変更 (旧版)"""
    doc = Document()
    doc.add_heading("IT サポート FAQ", level=0)
    doc.add_paragraph()
    
    doc.add_heading("Q. パスワードを忘れました", level=1)
    _add_para(doc, "A. IT部門（内線: 1234）に電話してリセット依頼をしてください。")
    
    path = AUTO_FIX_DIR / "IT_Support_FAQ.docx"
    doc.save(str(path))
    print(f"[OK] {path}")

def create_a4_it_contact_update():
    """A-4: 連絡先情報変更 (新版)"""
    doc = Document()
    doc.add_heading("IT部門 連絡先変更のお知らせ", level=0)
    doc.add_paragraph()
    
    doc.add_heading("新しい連絡方法", level=1)
    _add_para(doc, "内線1234は廃止されました。")
    _add_para(doc, "今後のお問い合わせは Slack #it-support チャンネルでお願いします。")
    
    path = AUTO_FIX_DIR / "IT_Contact_Update.docx"
    doc.save(str(path))
    print(f"[OK] {path}")

def create_a5_user_manual_v1_txt():
    """A-5: 制限値変更 (旧版) - TEXT形式"""
    content = """User Manual v1.0
Last Updated: 2024-07-01

## Chat Feature

Group chat can be created with up to 5 members maximum.
For larger groups, please use mailing lists instead.
"""
    path = AUTO_FIX_DIR / "User_Manual_v1.txt"
    path.write_text(content, encoding='utf-8')
    print(f"[OK] {path}")

def create_a5_system_update_v2_txt():
    """A-5: 制限値変更 (新版) - TEXT形式"""
    content = """System Update v2.0
Release Date: 2025-01-15

## Chat Feature Enhancement

Group chat member limit has been removed.
You can now create unlimited group chats with any number of members.
"""
    path = AUTO_FIX_DIR / "System_Update_v2.txt"
    path.write_text(content, encoding='utf-8')
    print(f"[OK] {path}")

def create_a6_employee_handbook_md():
    """A-6: ポリシー変更 (旧版) - MARKDOWN形式"""
    content = """# Employee Handbook 2024

## Remote Work Policy

- Remote work is allowed up to **2 days per week**
- Must be requested via email to your manager by the previous day
- Attendance tracking is required at start and end of shift
"""
    path = AUTO_FIX_DIR / "Employee_Handbook.md"
    path.write_text(content, encoding='utf-8')
    print(f"[OK] {path}")

def create_a6_hr_policy_2025_md():
    """A-6: ポリシー変更 (新版) - MARKDOWN形式"""
    content = """# HR Policy Update 2025

## Remote Work Policy Revision

- **Weekly limit removed**: Full remote work is now allowed
- **Application method changed**: Use HR Portal instead of email
- **Attendance tracking simplified**: Automatic logging, no manual input required
"""
    path = AUTO_FIX_DIR / "HR_Policy_Update_2025.md"
    path.write_text(content, encoding='utf-8')
    print(f"[OK] {path}")

# ============================================================================
# MANUAL ACTION REQUIRED (PDF placeholders)
# ============================================================================

def create_b1_setup_guide_v1_pdf():
    """B-1: PDF版バージョン混在 (v1) - docxで代用"""
    doc = Document()
    doc.add_heading("Setup Guide v1.0 (PDF)", level=0)
    _add_para(doc, "This is a PDF file placeholder. Cannot be auto-edited.")
    doc.add_paragraph()
    _add_para(doc, "Step 1: Download installer from legacy portal")
    _add_para(doc, "Step 2: Run setup.exe as Administrator")
    
    path = MANUAL_DIR / "Setup_Guide_v1.pdf.docx"  # .docx extension for demo
    doc.save(str(path))
    print(f"[OK] {path} (PDF placeholder)")

def create_b1_setup_guide_v2_pdf():
    """B-1: PDF版バージョン混在 (v2) - docxで代用"""
    doc = Document()
    doc.add_heading("Setup Guide v2.0 (PDF)", level=0)
    _add_para(doc, "This is a PDF file placeholder. Cannot be auto-edited.")
    doc.add_paragraph()
    _add_para(doc, "Step 1: Download installer from new cloud portal")
    _add_para(doc, "Step 2: Double-click to install (no admin required)")
    
    path = MANUAL_DIR / "Setup_Guide_v2.pdf.docx"
    doc.save(str(path))
    print(f"[OK] {path} (PDF placeholder)")

def create_b2_quick_start_pdf():
    """B-2: PDFスクリーンショット劣化 - docxで代用"""
    doc = Document()
    doc.add_heading("Quick Start Guide (PDF)", level=0)
    _add_para(doc, "This PDF contains outdated screenshots.")
    doc.add_paragraph()
    _add_para(doc, "[Screenshot: Old login screen with gear icon]")
    _add_para(doc, "Current UI has side menu instead.")
    
    path = MANUAL_DIR / "Quick_Start_v2.pdf.docx"
    doc.save(str(path))
    print(f"[OK] {path} (PDF placeholder)")

def create_b3_company_policy_pdf():
    """B-3: 混合フォーマット競合 - docxで代用"""
    doc = Document()
    doc.add_heading("Company Policy (PDF)", level=0)
    _add_para(doc, "Same content exists in both .docx and .pdf formats.")
    _add_para(doc, "This creates version ambiguity.")
    
    path = MANUAL_DIR / "Company_Policy.pdf.docx"
    doc.save(str(path))
    print(f"[OK] {path} (PDF placeholder)")

def create_b4_screenshot_png():
    """B-4: 画像ファイル単体 - テキストファイルで代用"""
    content = """[PNG IMAGE FILE]
Filename: old_dashboard_screenshot.png
Content: Screenshot of outdated dashboard UI (v2.0)
Issue: Current version is v3.0 with different layout
Action Required: Retake screenshot with v3.0 UI
"""
    path = MANUAL_DIR / "old_dashboard_screenshot.png.txt"
    path.write_text(content, encoding='utf-8')
    print(f"[OK] {path} (Image placeholder)")

# ============================================================================
# Main
# ============================================================================

if __name__ == "__main__":
    create_dirs()
    
    print("\n>> Generating AUTO-FIXABLE test documents...\n")
    create_a1_ui_guide_v2()
    create_a1_ui_changelog_v3()
    create_a2_onboarding_2024()
    create_a2_product_update_2025()
    create_a3_api_ref_v2()
    create_a3_api_migration_v3()
    create_a4_it_faq()
    create_a4_it_contact_update()
    create_a5_user_manual_v1_txt()
    create_a5_system_update_v2_txt()
    create_a6_employee_handbook_md()
    create_a6_hr_policy_2025_md()
    
    print("\n>> Generating MANUAL ACTION REQUIRED test documents...\n")
    create_b1_setup_guide_v1_pdf()
    create_b1_setup_guide_v2_pdf()
    create_b2_quick_start_pdf()
    create_b3_company_policy_pdf()
    create_b4_screenshot_png()
    
    print("\n>> Complete! Test documents created:")
    print(f"   Auto-Fix: {len(list(AUTO_FIX_DIR.glob('*')))} files")
    print(f"   Manual:   {len(list(MANUAL_DIR.glob('*')))} files")
    print(f"\n   Total: {len(list(AUTO_FIX_DIR.glob('*'))) + len(list(MANUAL_DIR.glob('*')))} test scenario files")
